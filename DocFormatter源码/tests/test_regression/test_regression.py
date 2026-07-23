"""输出一致性回归测试 — 确保每次输出的 docx 与基线一致.

原理:
1. 首次运行: 生成基线文件存入 tests/regression/baseline/
2. 后续运行: 与基线对比,任何差异都视为潜在回归
3. 差异判定: 对比关键内容（段落文本、表格数据、样式），而非二进制字节
   （因为 docx 是 ZIP 包，内部有随机字段如修改时间）
"""
import pytest
from pathlib import Path
from docx import Document as DocxDocument

from app.importers.md_importer import MarkdownImporter
from app.importers.txt_importer import TxtImporter
from app.templates import TemplateManager
from app.engine import Typesetter
from app.exporter import DocxExporter


BASELINE_DIR = Path(__file__).parent / "baseline"


def _extract_text_content(docx_path: Path) -> dict:
    """从 docx 提取可对比的文本内容."""
    docx = DocxDocument(str(docx_path))

    paragraphs = []
    for p in docx.paragraphs:
        paragraphs.append({
            "style": p.style.name if p.style else None,
            "text": p.text,
            "alignment": str(p.alignment) if p.alignment is not None else None,
        })

    tables = []
    for table in docx.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)

    return {
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "table_count": len(tables),
        "tables": tables,
        "section_count": len(docx.sections),
    }


def _compare_content(baseline: dict, current: dict) -> list:
    """对比两个内容字典，返回差异列表."""
    diffs = []

    if baseline["paragraph_count"] != current["paragraph_count"]:
        diffs.append(f"段落数不同: 基线 {baseline['paragraph_count']} vs 当前 {current['paragraph_count']}")

    if baseline["table_count"] != current["table_count"]:
        diffs.append(f"表格数不同: 基线 {baseline['table_count']} vs 当前 {current['table_count']}")

    # 检查段落文本
    for i, (bp, cp) in enumerate(zip(baseline["paragraphs"], current["paragraphs"])):
        if bp["text"] != cp["text"]:
            diffs.append(f"段落 {i+1} 文本不同: '{bp['text'][:30]}' vs '{cp['text'][:30]}'")
        if bp["style"] != cp["style"]:
            diffs.append(f"段落 {i+1} 样式不同: {bp['style']} vs {cp['style']}")

    # 检查表格内容
    for i, (bt, ct) in enumerate(zip(baseline["tables"], current["tables"])):
        if bt != ct:
            diffs.append(f"表格 {i+1} 内容不同")

    return diffs


@pytest.fixture
def ensure_baseline_dir():
    """确保基线目录存在."""
    BASELINE_DIR.mkdir(exist_ok=True)
    return BASELINE_DIR


class TestRegression:
    """回归测试."""

    @pytest.mark.p0
    def test_simple_markdown_regression(self, temp_md_file, temp_dir, ensure_baseline_dir):
        """简单 Markdown 的回归测试."""
        md = """# 文档标题

这是正文段落。

## 第一节

- 列表项一
- 列表项二

## 第二节

| 列A | 列B |
|------|------|
| 1 | 2 |
"""
        md_file = temp_md_file(md)
        baseline_path = ensure_baseline_dir / "simple_md.docx"

        # 生成当前输出
        doc = MarkdownImporter().import_file(md_file)
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        current_path = temp_dir / "current.docx"
        DocxExporter().export(doc, current_path)

        current_content = _extract_text_content(current_path)

        if not baseline_path.exists():
            # 首次运行: 生成基线
            baseline_path.write_bytes(current_path.read_bytes())
            pytest.skip(f"基线文件已生成: {baseline_path}")

        # 对比基线
        baseline_content = _extract_text_content(baseline_path)
        diffs = _compare_content(baseline_content, current_content)

        assert not diffs, "与基线存在差异:\n" + "\n".join(diffs)

    @pytest.mark.p0
    def test_complex_markdown_regression(self, temp_md_file, temp_dir, ensure_baseline_dir):
        """复杂 Markdown 的回归测试."""
        md = """# 总论

这是包含**加粗**、*斜体*、`代码`的段落。

## 列表

- 无序一
- 无序二

1. 有序一
2. 有序二

## 表格

| 姓名 | 年龄 | 城市 |
|------|------|------|
| 张三 | 25 | 北京 |
| 李四 | 30 | 上海 |

## 引用

> 引用内容

## 代码

```python
def hello():
    print("world")
```
"""
        md_file = temp_md_file(md)
        baseline_path = ensure_baseline_dir / "complex_md.docx"

        doc = MarkdownImporter().import_file(md_file)
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        current_path = temp_dir / "current.docx"
        DocxExporter().export(doc, current_path)

        current_content = _extract_text_content(current_path)

        if not baseline_path.exists():
            baseline_path.write_bytes(current_path.read_bytes())
            pytest.skip(f"基线文件已生成: {baseline_path}")

        baseline_content = _extract_text_content(baseline_path)
        diffs = _compare_content(baseline_content, current_content)

        assert not diffs, "与基线存在差异:\n" + "\n".join(diffs)

    @pytest.mark.p0
    def test_chinese_content_regression(self, temp_md_file, temp_dir, ensure_baseline_dir):
        """中文内容的回归测试."""
        md = """# 中文文档测试

这是一个**中文文档**，包含各种中文标点符号：，。！？；：""''【】《》。

## 第一节

中文正文段落，需要正确的首行缩进和两端对齐。

## 第二节

更多内容测试。"""

        md_file = temp_md_file(md)
        baseline_path = ensure_baseline_dir / "chinese.docx"

        doc = MarkdownImporter().import_file(md_file)
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        current_path = temp_dir / "current.docx"
        DocxExporter().export(doc, current_path)

        current_content = _extract_text_content(current_path)

        if not baseline_path.exists():
            baseline_path.write_bytes(current_path.read_bytes())
            pytest.skip(f"基线文件已生成: {baseline_path}")

        baseline_content = _extract_text_content(baseline_path)
        diffs = _compare_content(baseline_content, current_content)

        assert not diffs, "与基线存在差异:\n" + "\n".join(diffs)

    @pytest.mark.p0
    def test_real_docx_regression(self, temp_dir, ensure_baseline_dir):
        """使用真实 docx 文件的回归测试."""
        real_file = Path("/Users/mac/Documents/量化/office/全省第一台工伤预防大篷车亮相.docx")
        if not real_file.exists():
            pytest.skip("真实 docx 文件不存在")

        from app.importers.docx_importer import DocxImporter

        baseline_path = ensure_baseline_dir / "real_docx.docx"

        # 导入 → 排版 → 导出
        doc = DocxImporter().import_file(real_file)
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        current_path = temp_dir / "current.docx"
        DocxExporter().export(doc, current_path)

        current_content = _extract_text_content(current_path)

        if not baseline_path.exists():
            baseline_path.write_bytes(current_path.read_bytes())
            pytest.skip(f"基线文件已生成: {baseline_path}")

        baseline_content = _extract_text_content(baseline_path)
        diffs = _compare_content(baseline_content, current_content)

        assert not diffs, "与基线存在差异:\n" + "\n".join(diffs)


class TestContentPreservation:
    """内容保留测试 — 确保往返过程中内容不丢失."""

    @pytest.mark.p0
    def test_all_paragraphs_preserved(self, temp_md_file, temp_dir):
        """所有段落在往返后保留."""
        md = "\n\n".join([f"段落 {i}" for i in range(50)])
        md_file = temp_md_file(md)

        doc = MarkdownImporter().import_file(md_file)
        original_texts = [e.get_text() for e in doc.body if hasattr(e, 'get_text')]

        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)

        output = temp_dir / "output.docx"
        DocxExporter().export(doc, output)

        # 读回
        docx = DocxDocument(str(output))
        roundtrip_texts = [p.text for p in docx.paragraphs if p.text.strip()]

        # 验证所有原始文本都在
        for original in original_texts:
            assert any(original in rt for rt in roundtrip_texts), \
                f"丢失段落: {original}"

    @pytest.mark.p0
    def test_headings_preserved_with_level(self, temp_md_file, temp_dir):
        """标题层级在往返后保留."""
        md = """# H1
## H2
### H3
#### H4"""
        md_file = temp_md_file(md)

        doc = MarkdownImporter().import_file(md_file)
        original_headings = [(h.get_text(), h.level.value) for h in doc.get_headings()]

        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        output = temp_dir / "output.docx"
        DocxExporter().export(doc, output)

        # 读回验证标题
        docx = DocxDocument(str(output))
        headings = []
        for p in docx.paragraphs:
            if p.style.name.startswith("Heading"):
                level = int(p.style.name.split()[-1])
                headings.append((p.text, level))

        assert len(headings) == len(original_headings)
        for (orig_text, orig_level), (rt_text, rt_level) in zip(original_headings, headings):
            assert orig_text in rt_text, f"标题文本不匹配: {orig_text} vs {rt_text}"
            assert orig_level == rt_level, f"标题层级不匹配: H{orig_level} vs H{rt_level}"

    @pytest.mark.p0
    def test_table_data_preserved(self, temp_md_file, temp_dir):
        """表格数据在往返后保留."""
        md = """| 姓名 | 年龄 | 城市 |
|------|------|------|
| 张三 | 25 | 北京 |
| 李四 | 30 | 上海 |"""
        md_file = temp_md_file(md)

        doc = MarkdownImporter().import_file(md_file)
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)
        output = temp_dir / "output.docx"
        DocxExporter().export(doc, output)

        docx = DocxDocument(str(output))
        assert len(docx.tables) == 1

        table = docx.tables[0]
        assert table.cell(0, 0).text.strip() == "姓名"
        assert table.cell(1, 1).text.strip() == "25"
        assert table.cell(2, 2).text.strip() == "上海"
