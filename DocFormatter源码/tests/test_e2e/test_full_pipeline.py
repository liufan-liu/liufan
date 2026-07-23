"""端到端测试 — 验证完整业务流程."""
import pytest
from pathlib import Path
from docx import Document as DocxDocument

from app.importers.md_importer import MarkdownImporter
from app.importers.txt_importer import TxtImporter
from app.importers.base import ImporterRegistry
from app.templates import TemplateManager
from app.engine import Typesetter
from app.engine.toc_generator import TOCGenerator
from app.exporter import DocxExporter


class TestEndToEnd:
    """端到端集成测试."""

    @pytest.mark.p0
    def test_txt_full_pipeline(self, temp_txt_file, temp_docx_path):
        """EE-001: TXT 全链路."""
        # 1. 写入测试 TXT
        f = temp_txt_file("""第一章 项目简介

这是项目简介段落。

1.1 背景

背景内容描述。

第二章 技术方案

技术方案详情。""")

        # 2. 导入
        doc = TxtImporter().import_file(f)
        assert doc.element_count() > 0

        # 3. 排版
        tm = TemplateManager()
        tpl = tm.get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)

        # 4. 导出
        DocxExporter().export(doc, temp_docx_path)
        assert temp_docx_path.exists()

        # 5. 验证
        docx = DocxDocument(str(temp_docx_path))
        para_texts = [p.text for p in docx.paragraphs]
        assert any("项目简介" in t for t in para_texts)
        assert any("技术方案" in t for t in para_texts)

        # 验证标题样式
        headings = [p for p in docx.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 3

    @pytest.mark.p0
    def test_markdown_full_pipeline(self, temp_md_file, temp_docx_path):
        """EE-002: Markdown 全链路."""
        md = """# 文档标题

这是**加粗**正文内容。

## 第一节

- 列表项一
- 列表项二

## 第二节

| 列A | 列B |
|------|------|
| 1 | 2 |

## 第三节

> 引用文字

```python
print("hello")
```
"""
        f = temp_md_file(md)

        # 导入
        doc = MarkdownImporter().import_file(f)
        assert doc.element_count() >= 8

        # 排版
        tm = TemplateManager()
        tpl = tm.get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)

        # 导出
        DocxExporter().export(doc, temp_docx_path)
        assert temp_docx_path.exists()

        # 验证
        docx = DocxDocument(str(temp_docx_path))
        para_texts = [p.text for p in docx.paragraphs]
        assert any("文档标题" in t for t in para_texts)
        assert any("加粗" in t for t in para_texts)

        # 验证表格
        assert len(docx.tables) >= 1

    @pytest.mark.p0
    def test_complex_markdown_no_loss(self, temp_md_file, temp_docx_path):
        """EE-003: 复杂 Markdown 无元素丢失."""
        md = """# 总论

## 1.1 背景

普通段落**粗体**和*斜体*。

## 1.2 列表

- 无序一
- 无序二

1. 有序一
2. 有序二

## 1.3 表格

| 项目 | 数值 |
|------|------|
| A | 100 |
| B | 200 |

## 1.4 代码

```python
def hello():
    print("world")
```

## 1.5 引用

> 这是一段引用文字

---

# 结论

结束。
"""
        f = temp_md_file(md)

        # 导入
        doc = MarkdownImporter().import_file(f)
        input_types = {type(e).__name__ for e in doc.body}

        # 排版
        tpl = TemplateManager().get_template("通用文档")
        doc = Typesetter(tpl).apply(doc)

        # 导出并读回
        DocxExporter().export(doc, temp_docx_path)
        docx = DocxDocument(str(temp_docx_path))

        # 验证关键内容
        all_text = "\n".join(p.text for p in docx.paragraphs)
        assert "总论" in all_text
        assert "背景" in all_text
        assert "粗体" in all_text
        assert "结论" in all_text

        # 验证表格
        assert len(docx.tables) >= 1

        # 验证标题数量
        headings = [p for p in docx.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 5

    @pytest.mark.p0
    def test_template_switch(self, temp_md_file, temp_dir):
        """EE-004: 切换模板产生不同样式."""
        f = temp_md_file("# 标题\n\n正文内容。")

        # 用通用文档模板
        doc1 = MarkdownImporter().import_file(f)
        tpl1 = TemplateManager().get_template("通用文档")
        doc1 = Typesetter(tpl1).apply(doc1)
        path1 = temp_dir / "output1.docx"
        DocxExporter().export(doc1, path1)

        # 用自定义模板（不同字体/字号）
        doc2 = MarkdownImporter().import_file(f)
        tpl2 = {
            "name": "其他模板",
            "page": {"width": 210, "height": 297,
                     "margin_top": 40, "margin_bottom": 40,
                     "margin_left": 40, "margin_right": 40},
            "styles": {
                "Heading 1": {
                    "run": {"font_name": "Arial", "font_size": 24, "bold": True},
                    "paragraph": {"alignment": "left"},
                },
                "Normal": {
                    "run": {"font_name": "Georgia", "font_size": 14},
                    "paragraph": {"alignment": "justify", "line_spacing": 2.0},
                },
            },
        }
        doc2 = Typesetter(tpl2).apply(doc2)
        path2 = temp_dir / "output2.docx"
        DocxExporter().export(doc2, path2)

        # 比较两个输出
        docx1 = DocxDocument(str(path1))
        docx2 = DocxDocument(str(path2))

        # 页面设置不同
        assert docx1.sections[0].top_margin.mm != docx2.sections[0].top_margin.mm

        # 正文字体不同
        para1 = next(p for p in docx1.paragraphs if "正文" in p.text)
        para2 = next(p for p in docx2.paragraphs if "正文" in p.text)
        assert para1.runs[0].font.size != para2.runs[0].font.size

    @pytest.mark.p0
    def test_chinese_content_no_garble(self, temp_md_file, temp_docx_path):
        """EE-005: 中文内容无乱码."""
        md = """# 中文标题测试

这是一段包含中文、English、日本語、한국어 的多语言文本。

## 特殊字符

符号：①②③ ④⑤⑥ 【】《》「」『』
标点：，。！？；：""''
数学：± × ÷ ≤ ≥ ≠ ≈
货币：$ € £ ¥ ₽

## 长段落

这是一段很长的中文文本，用来测试排版引擎在处理长段落时的表现。DocFormatter 需要能够正确处理各种中文字符，包括标点符号、特殊符号以及不同字体下的显示效果。我们希望导出的 Word 文档能够在不同平台上正确显示，不出现乱码或排版错乱的情况。
"""
        f = temp_md_file(md)

        doc = MarkdownImporter().import_file(f)
        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, temp_docx_path)

        # 读回验证
        docx = DocxDocument(str(temp_docx_path))
        all_text = "\n".join(p.text for p in docx.paragraphs)

        # 验证各种字符
        assert "中文标题" in all_text
        assert "English" in all_text
        assert "日本語" in all_text
        assert "①②③" in all_text
        assert "，。！？" in all_text
        assert "± × ÷" in all_text
        assert "长段落" in all_text


class TestEdgeCases:
    """边界条件测试."""

    @pytest.mark.p1
    def test_empty_markdown(self, temp_md_file, temp_docx_path):
        """空 Markdown 文件抛出友好错误（按规范，空文件属于异常输入）."""
        from app.importers.base import ImportError as ImportErr
        f = temp_md_file("")
        with pytest.raises(ImportErr, match="文件为空"):
            MarkdownImporter().import_file(f)

    @pytest.mark.p1
    def test_single_character(self, temp_md_file, temp_docx_path):
        """单字符文件."""
        f = temp_md_file("A")
        doc = MarkdownImporter().import_file(f)
        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, temp_docx_path)
        assert temp_docx_path.exists()

    @pytest.mark.p1
    def test_special_characters(self, temp_md_file, temp_docx_path):
        """特殊字符处理."""
        f = temp_md_file("Emoji: 🎉🎊\nMath: ∑∫∞\nSymbols: ©®™")
        doc = MarkdownImporter().import_file(f)
        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, temp_docx_path)

        docx = DocxDocument(str(temp_docx_path))
        all_text = "\n".join(p.text for p in docx.paragraphs)
        # 基本字符应该保留
        assert "Math:" in all_text or "Symbol:" in all_text

    @pytest.mark.p1
    def test_very_long_paragraph(self, temp_md_file, temp_docx_path):
        """超长段落（10000 字符）."""
        long_text = "测试" * 5000  # 10000 字符
        f = temp_md_file(f"# 标题\n\n{long_text}")
        doc = MarkdownImporter().import_file(f)
        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, temp_docx_path)
        assert temp_docx_path.exists()
        assert temp_docx_path.stat().st_size > 0

    @pytest.mark.p1
    def test_many_headings(self, temp_md_file, temp_docx_path):
        """大量标题（100+）."""
        md_lines = []
        for i in range(100):
            md_lines.append(f"## 第 {i+1} 节")
            md_lines.append(f"这是第 {i+1} 节的内容。")
        f = temp_md_file("\n\n".join(md_lines))

        doc = MarkdownImporter().import_file(f)
        assert len(doc.get_headings()) == 100

        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, temp_docx_path)

        docx = DocxDocument(str(temp_docx_path))
        headings = [p for p in docx.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 100

    @pytest.mark.p1
    def test_chinese_path(self, temp_dir, temp_md_file):
        """中文路径."""
        chinese_dir = temp_dir / "中文目录" / "测试"
        chinese_dir.mkdir(parents=True)
        output = chinese_dir / "输出文件.docx"

        f = temp_md_file("# 标题\n\n内容")
        doc = MarkdownImporter().import_file(f)
        doc = Typesetter(TemplateManager().get_template("通用文档")).apply(doc)
        DocxExporter().export(doc, output)
        assert output.exists()
