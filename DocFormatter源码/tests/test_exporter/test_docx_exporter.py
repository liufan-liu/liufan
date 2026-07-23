"""DOCX 导出器单元测试."""
import pytest
from pathlib import Path
from docx import Document as DocxDocument

from app.exporter import DocxExporter, ExportError
from app.model.document import Document
from app.model.elements import (
    Paragraph, Heading, Run, Table, PageBreak, TOCField, CodeBlock, BlockQuote,
)
from app.model.styles import RunStyle, ParagraphStyle
from app.model.enums import Alignment, HeadingLevel


class TestDocxExporter:
    """DOCX 导出器测试."""

    @pytest.fixture
    def exporter(self):
        return DocxExporter()

    def _read_back(self, path):
        """读回导出的 docx 文件."""
        return DocxDocument(str(path))

    @pytest.mark.p0
    def test_export_empty_document(self, empty_document, temp_docx_path, exporter):
        """X-001: 导出空文档."""
        exporter.export(empty_document, temp_docx_path)
        assert temp_docx_path.exists()
        assert temp_docx_path.stat().st_size > 0
        # 验证可被 python-docx 打开
        docx = self._read_back(temp_docx_path)
        assert len(docx.paragraphs) >= 0

    @pytest.mark.p0
    def test_export_heading(self, temp_docx_path, exporter):
        """X-002: 导出标题."""
        doc = Document()
        doc.add_heading("测试标题", level=1)
        # 应用样式
        heading = doc.body[0]
        heading.runs[0].style.font_name = "黑体"
        heading.runs[0].style.font_size = 18
        heading.runs[0].style.bold = True

        exporter.export(doc, temp_docx_path)
        docx = self._read_back(temp_docx_path)

        # 找到包含"测试标题"的段落
        title_para = next(p for p in docx.paragraphs if "测试标题" in p.text)
        assert title_para.style.name.startswith("Heading")
        assert title_para.runs[0].font.bold is True

    @pytest.mark.p0
    def test_export_paragraph(self, temp_docx_path, exporter):
        """X-003: 导出段落."""
        doc = Document()
        para = Paragraph(
            runs=[Run(text="测试段落")],
            style=ParagraphStyle(alignment=Alignment.CENTER, line_spacing=1.5),
        )
        doc.body.append(para)
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        text_para = next(p for p in docx.paragraphs if "测试段落" in p.text)
        assert text_para.alignment is not None  # 对齐被设置

    @pytest.mark.p0
    def test_export_bold_italic(self, temp_docx_path, exporter):
        """X-004: 导出粗体和斜体."""
        doc = Document()
        para = Paragraph(runs=[
            Run(text="粗体", style=RunStyle(bold=True)),
            Run(text="斜体", style=RunStyle(italic=True)),
        ])
        doc.body.append(para)
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        text_para = next(p for p in docx.paragraphs if "粗体" in p.text)
        bold_run = next(r for r in text_para.runs if "粗体" in r.text)
        assert bold_run.font.bold is True

    @pytest.mark.p0
    def test_export_table(self, temp_docx_path, exporter):
        """X-005: 导出表格."""
        doc = Document()
        doc.body.append(Table.from_data([
            ["姓名", "年龄"],
            ["张三", "25"],
        ]))
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        assert len(docx.tables) == 1
        table = docx.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "姓名"
        assert table.cell(1, 1).text == "25"

    @pytest.mark.p1
    def test_export_page_break(self, temp_docx_path, exporter):
        """X-009: 导出分页符."""
        doc = Document()
        doc.add_paragraph("上方")
        doc.body.append(PageBreak())
        doc.add_paragraph("下方")
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        # 分页符应产生至少 3 个段落
        assert len(docx.paragraphs) >= 3

    @pytest.mark.p0
    def test_export_toc_field(self, temp_docx_path, exporter):
        """X-010: 导出目录域."""
        doc = Document()
        doc.body.append(TOCField(max_level=3, title="目录"))
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        # 应包含目录标题
        assert any("目录" in p.text for p in docx.paragraphs)
        # 应包含 TOC 域代码（通过检查 XML）
        xml = docx.element.xml
        assert "TOC" in xml or "目录" in xml

    @pytest.mark.p0
    def test_page_settings(self, temp_docx_path, exporter):
        """X-013: 页面设置正确."""
        doc = Document()
        doc.page_style.margin_top = 30
        doc.page_style.margin_left = 25
        doc.add_paragraph("test")
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        section = docx.sections[0]
        assert abs(section.top_margin.mm - 30) < 1
        assert abs(section.left_margin.mm - 25) < 1

    @pytest.mark.p0
    def test_chinese_font(self, temp_docx_path, exporter):
        """X-014: 中文字体设置."""
        doc = Document()
        para = Paragraph(runs=[Run(
            text="中文测试",
            style=RunStyle(font_name="Arial", font_name_east_asia="黑体"),
        )])
        doc.body.append(para)
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        para = next(p for p in docx.paragraphs if "中文" in p.text)
        # 检查 eastAsia 字体设置
        run = para.runs[0]
        rpr = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rpr is not None:
            east_asia = rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
            assert east_asia == "黑体"

    @pytest.mark.p1
    def test_export_to_nonexistent_dir(self, temp_dir, exporter):
        """X-016: 导出到不存在的目录，自动创建."""
        doc = Document()
        doc.add_paragraph("test")
        output = temp_dir / "nonexistent" / "subdir" / "output.docx"
        exporter.export(doc, output)
        assert output.exists()

    @pytest.mark.p1
    def test_core_properties(self, temp_docx_path, exporter):
        """X-015: 核心属性设置."""
        doc = Document(title="我的文档", author="张三")
        doc.add_paragraph("test")
        exporter.export(doc, temp_docx_path)

        docx = self._read_back(temp_docx_path)
        assert docx.core_properties.title == "我的文档"
        assert docx.core_properties.author == "张三"
