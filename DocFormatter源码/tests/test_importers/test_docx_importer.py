"""DOCX 导入器单元测试."""
import pytest
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.importers.docx_importer import DocxImporter
from app.importers.base import ImportError
from app.model.elements import Paragraph, Heading, Table, PageBreak
from app.model.enums import HeadingLevel, Alignment


@pytest.fixture
def make_docx(temp_dir):
    """创建测试用 docx 文件的工厂 fixture."""
    def _make(filename="test.docx", setup_fn=None):
        docx = DocxDocument()
        path = temp_dir / filename
        if setup_fn:
            setup_fn(docx)
        docx.save(str(path))
        return path
    return _make


class TestDocxImporter:
    """DOCX 导入器测试."""

    @pytest.mark.p0
    def test_can_import_docx(self, docx_importer):
        """识别 .docx 扩展名."""
        assert docx_importer.can_import(Path("test.docx")) is True
        assert docx_importer.can_import(Path("test.DOCX")) is True
        assert docx_importer.can_import(Path("test.txt")) is False

    @pytest.mark.p0
    def test_import_simple_paragraphs(self, make_docx, docx_importer):
        """导入含多个段落的 docx."""
        def setup(docx):
            docx.add_paragraph("第一段内容。")
            docx.add_paragraph("第二段内容。")
            docx.add_paragraph("第三段内容。")

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        paras = [e for e in doc.body if isinstance(e, Paragraph)]
        assert len(paras) == 3
        assert paras[0].get_text() == "第一段内容。"
        assert paras[1].get_text() == "第二段内容。"
        assert paras[2].get_text() == "第三段内容。"

    @pytest.mark.p0
    def test_import_headings(self, make_docx, docx_importer):
        """导入标题（按 Word 标题样式识别）."""
        def setup(docx):
            docx.add_heading("一级标题", level=1)
            docx.add_paragraph("正文。")
            docx.add_heading("二级标题", level=2)

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        headings = doc.get_headings()
        assert len(headings) == 2
        assert headings[0].level == HeadingLevel.H1
        assert headings[0].get_text() == "一级标题"
        assert headings[1].level == HeadingLevel.H2

    @pytest.mark.p0
    def test_import_run_styles(self, make_docx, docx_importer):
        """导入 Run 级样式（字体、粗体、斜体）."""
        def setup(docx):
            para = docx.add_paragraph()
            run1 = para.add_run("普通文本")
            run2 = para.add_run("加粗文本")
            run2.bold = True
            run3 = para.add_run("斜体文本")
            run3.italic = True

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        para = doc.body[0]
        assert len(para.runs) == 3
        assert para.runs[1].style.bold is True
        assert para.runs[2].style.italic is True

    @pytest.mark.p0
    def test_import_font_name(self, make_docx, docx_importer):
        """导入字体名称."""
        def setup(docx):
            para = docx.add_paragraph()
            run = para.add_run("宋体文字")
            run.font.name = "宋体"

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        run = doc.body[0].runs[0]
        assert run.style.font_name == "宋体"

    @pytest.mark.p0
    def test_import_font_size(self, make_docx, docx_importer):
        """导入字号."""
        def setup(docx):
            para = docx.add_paragraph()
            run = para.add_run("16pt 文字")
            run.font.size = Pt(16)

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        run = doc.body[0].runs[0]
        assert run.style.font_size == 16

    @pytest.mark.p0
    def test_import_table(self, make_docx, docx_importer):
        """导入表格."""
        def setup(docx):
            table = docx.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "姓名"
            table.cell(0, 1).text = "年龄"
            table.cell(1, 0).text = "张三"
            table.cell(1, 1).text = "25"
            table.cell(2, 0).text = "李四"
            table.cell(2, 1).text = "30"

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        tables = [e for e in doc.body if isinstance(e, Table)]
        assert len(tables) == 1
        assert tables[0].rows == 3
        assert tables[0].cols == 2
        assert tables[0].cells[0][0].text == "姓名"
        assert tables[0].cells[1][1].text == "25"

    @pytest.mark.p0
    def test_import_page_settings(self, make_docx, docx_importer):
        """导入页面设置."""
        def setup(docx):
            section = docx.sections[0]
            section.top_margin = Cm(3)
            section.left_margin = Cm(2.5)

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        assert abs(doc.page_style.margin_top - 30) < 0.5  # 3cm = 30mm
        assert abs(doc.page_style.margin_left - 25) < 0.5  # 2.5cm = 25mm

    @pytest.mark.p1
    def test_import_paragraph_alignment(self, make_docx, docx_importer):
        """导入段落对齐方式."""
        def setup(docx):
            para = docx.add_paragraph("居中对齐")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        para = doc.body[0]
        assert para.style.alignment == Alignment.CENTER

    @pytest.mark.p1
    def test_import_document_title(self, make_docx, docx_importer):
        """导入文档标题（从核心属性）."""
        def setup(docx):
            docx.core_properties.title = "我的文档标题"
            docx.core_properties.author = "张三"
            docx.add_paragraph("正文")

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        assert doc.title == "我的文档标题"
        assert doc.author == "张三"

    @pytest.mark.p0
    def test_import_empty_docx(self, make_docx, docx_importer):
        """导入空 docx 文件."""
        path = make_docx()
        doc = docx_importer.import_file(path)
        assert doc.element_count() == 0

    @pytest.mark.p0
    def test_import_nonexistent_raises(self, docx_importer):
        """不存在文件抛出 ImportError."""
        with pytest.raises(ImportError, match="不存在"):
            docx_importer.import_file(Path("/nonexistent/file.docx"))

    @pytest.mark.p0
    def test_import_empty_file_raises(self, temp_dir, docx_importer):
        """空文件抛出 ImportError."""
        empty = temp_dir / "empty.docx"
        empty.write_bytes(b"")
        with pytest.raises(ImportError, match="为空"):
            docx_importer.import_file(empty)

    @pytest.mark.p1
    def test_import_mixed_content(self, make_docx, docx_importer):
        """导入混合内容（标题 + 段落 + 表格）."""
        def setup(docx):
            docx.add_heading("标题", level=1)
            docx.add_paragraph("正文段落。")
            table = docx.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(1, 0).text = "C"
            table.cell(1, 1).text = "D"
            docx.add_paragraph("结尾段落。")

        path = make_docx(setup_fn=setup)
        doc = docx_importer.import_file(path)

        assert len(doc.get_headings()) == 1
        tables = [e for e in doc.body if isinstance(e, Table)]
        assert len(tables) == 1
        paras = [e for e in doc.body if type(e) is Paragraph]
        assert len(paras) == 2

    @pytest.mark.p1
    def test_roundtrip_txt_to_docx_to_doc(self, temp_dir, txt_importer, docx_importer):
        """TXT → DOCX → Document 往返测试."""
        from app.engine import Typesetter
        from app.templates import TemplateManager
        from app.exporter import DocxExporter

        # 1. 写 TXT
        txt_file = temp_dir / "test.txt"
        txt_file.write_text("第一章 简介\n\n这是正文。\n\n第二章 总结\n\n总结内容。", encoding="utf-8")

        # 2. 导入 TXT
        doc1 = txt_importer.import_file(txt_file)
        assert doc1.element_count() > 0

        # 3. 排版 + 导出为 DOCX
        tpl = TemplateManager().get_template("通用文档")
        doc1 = Typesetter(tpl).apply(doc1)
        docx_file = temp_dir / "output.docx"
        DocxExporter().export(doc1, docx_file)

        # 4. 导入 DOCX
        doc2 = docx_importer.import_file(docx_file)
        assert doc2.element_count() > 0

        # 5. 验证内容保留
        texts = [e.get_text() for e in doc2.body if hasattr(e, 'get_text')]
        full_text = "\n".join(texts)
        assert "简介" in full_text
        assert "正文" in full_text
        assert "总结" in full_text


class TestDocxImporterWithRealFiles:
    """使用真实 docx 文件的集成测试."""

    @pytest.mark.p0
    def test_import_real_docx(self):
        """导入用户目录下的真实 docx 文件."""
        real_file = Path("/Users/mac/Documents/量化/office/全省第一台工伤预防大篷车亮相.docx")
        if not real_file.exists():
            pytest.skip("真实 docx 文件不存在")

        doc = DocxImporter().import_file(real_file)
        assert doc.element_count() > 0
        assert doc.title == "全省第一台工伤预防大篷车亮相"

        # 验证能提取文本
        texts = [e.get_text() for e in doc.body if hasattr(e, 'get_text')]
        full_text = "\n".join(texts)
        assert len(full_text) > 0
        # 应该能读到一些中文字符
        assert any("车" in t or "工伤" in t for t in texts)

    @pytest.mark.p1
    def test_real_docx_roundtrip(self, temp_dir):
        """真实 docx 文件的往返测试：导入 → 排版 → 导出 → 再导入."""
        real_file = Path("/Users/mac/Documents/量化/office/全省第一台工伤预防大篷车亮相.docx")
        if not real_file.exists():
            pytest.skip("真实 docx 文件不存在")

        from app.engine import Typesetter
        from app.templates import TemplateManager
        from app.exporter import DocxExporter

        # 1. 导入
        doc1 = DocxImporter().import_file(real_file)
        original_count = doc1.element_count()

        # 2. 排版
        tpl = TemplateManager().get_template("通用文档")
        doc1 = Typesetter(tpl).apply(doc1)

        # 3. 导出
        output = temp_dir / "roundtrip.docx"
        DocxExporter().export(doc1, output)

        # 4. 再导入
        doc2 = DocxImporter().import_file(output)

        # 验证内容保留
        texts1 = [e.get_text() for e in doc1.body if hasattr(e, 'get_text')]
        texts2 = [e.get_text() for e in doc2.body if hasattr(e, 'get_text')]
        text1 = "\n".join(texts1)
        text2 = "\n".join(texts2)

        # 主要文本内容应大致保留
        assert len(text2) > len(text1) * 0.5  # 至少保留 50%
