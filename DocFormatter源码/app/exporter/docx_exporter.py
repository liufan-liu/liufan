"""DOCX 导出器 — 将 Document 模型写入 .docx 文件.

使用 python-docx 库生成 Word 文档，支持：
- 页面设置（纸张尺寸、边距）
- 样式应用（标题、段落、表格、列表、代码块）
- 图片嵌入
- 页眉页脚
- 目录域（TOC field）
- 分页符
"""
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.model.document import Document, HeaderFooter
from app.model.elements import (
    Paragraph, Heading, Table, TableCell, Image, PageBreak,
    TOCField, Run, DocumentList, ListItem, CodeBlock, BlockQuote,
)
from app.model.styles import RunStyle, ParagraphStyle, PageStyle
from app.model.enums import Alignment, HeadingLevel


class ExportError(Exception):
    """导出失败异常."""
    pass


class DocxExporter:
    """DOCX 导出器."""

    ALIGNMENT_MAP = {
        Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def export(self, doc: Document, output_path: Path) -> None:
        """将 Document 模型导出为 .docx 文件.

        Args:
            doc: 排版后的文档模型
            output_path: 输出文件路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        docx = DocxDocument()

        try:
            # 1. 设置核心属性
            self._apply_core_properties(docx, doc)

            # 2. 设置页面
            self._apply_page_settings(docx, doc)

            # 3. 设置页眉页脚
            if doc.header or doc.footer:
                self._apply_header_footer(docx, doc)

            # 4. 写入文档体
            for element in doc.body:
                self._write_element(docx, element, doc)

            # 5. 保存
            docx.save(str(output_path))

        except Exception as e:
            raise ExportError(f"导出失败: {e}") from e

    def _apply_core_properties(self, docx, doc: Document) -> None:
        """设置文档核心属性（标题、作者等）."""
        props = docx.core_properties
        if doc.title:
            props.title = doc.title
        if doc.author:
            props.author = doc.author
        if doc.subject:
            props.subject = doc.subject
        if doc.keywords:
            props.keywords = ",".join(doc.keywords)

    def _apply_page_settings(self, docx, doc: Document) -> None:
        """设置页面尺寸和边距."""
        ps = doc.page_style
        section = docx.sections[0]

        # 纸张尺寸 (mm → EMU)
        section.page_width = Cm(ps.width / 10)
        section.page_height = Cm(ps.height / 10)

        # 方向
        if ps.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT

        # 边距 (mm → cm → EMU)
        section.top_margin = Cm(ps.margin_top / 10)
        section.bottom_margin = Cm(ps.margin_bottom / 10)
        section.left_margin = Cm(ps.margin_left / 10)
        section.right_margin = Cm(ps.margin_right / 10)

        # 页眉页脚距离
        section.header_distance = Cm(ps.header_distance / 10)
        section.footer_distance = Cm(ps.footer_distance / 10)

    def _apply_header_footer(self, docx, doc: Document) -> None:
        """设置页眉页脚."""
        section = docx.sections[0]

        if doc.header and not doc.header.is_empty():
            header = section.header
            header.is_linked_to_previous = False
            self._write_header_footer_content(header, doc.header)

        if doc.footer and not doc.footer.is_empty():
            footer = section.footer
            footer.is_linked_to_previous = False
            self._write_header_footer_content(footer, doc.footer, is_footer=True)

    def _write_header_footer_content(self, hf_obj, data: HeaderFooter, is_footer: bool = False) -> None:
        """写入页眉/页脚内容."""
        # 清除默认段落
        for p in hf_obj.paragraphs:
            p.clear()

        para = hf_obj.paragraphs[0] if hf_obj.paragraphs else hf_obj.add_paragraph()

        # 左
        if data.left:
            para.add_run(data.left)

        # 中
        if data.center:
            # 用制表符实现左-中-右布局
            para.add_run("\t")
            para.add_run(data.center)

        # 右
        if data.right:
            para.add_run("\t\t")
            para.add_run(data.right)

        # 页码
        if data.include_page_number:
            if not data.center and not data.right:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_page_number_field(para)

        # 应用字体样式
        if hasattr(data, "_run_style"):
            rs = data._run_style
            for run in para.runs:
                if rs.font_name:
                    run.font.name = rs.font_name
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), rs.font_name_east_asia or rs.font_name)
                if rs.font_size:
                    run.font.size = Pt(rs.font_size)

    def _add_page_number_field(self, paragraph, font_name=None, font_size=None) -> None:
        """在段落中添加页码域代码."""
        run = paragraph.add_run()
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if font_size:
            run.font.size = Pt(font_size)

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

    def _write_element(self, docx, element, doc: Document) -> None:
        """写入单个元素到 docx."""
        if isinstance(element, Heading):
            self._write_heading(docx, element)
        elif isinstance(element, Paragraph):
            self._write_paragraph(docx, element)
        elif isinstance(element, Table):
            self._write_table(docx, element)
        elif isinstance(element, Image):
            self._write_image(docx, element)
        elif isinstance(element, PageBreak):
            docx.add_page_break()
        elif isinstance(element, TOCField):
            self._write_toc_field(docx, element)
        elif isinstance(element, DocumentList):
            self._write_list(docx, element)
        elif isinstance(element, CodeBlock):
            self._write_code_block(docx, element)
        elif isinstance(element, BlockQuote):
            self._write_blockquote(docx, element)

    def _write_heading(self, docx, heading: Heading) -> None:
        """写入标题."""
        p = docx.add_heading(level=heading.level.value)
        # 清除默认空 run
        p.clear()

        # 应用段落样式
        self._apply_paragraph_style(p, heading.style, heading.runs)

        # 写入 runs
        for run_data in heading.runs:
            run = p.add_run(run_data.text)
            self._apply_run_style(run, run_data.style)

    def _write_paragraph(self, docx, para: Paragraph) -> None:
        """写入普通段落."""
        p = docx.add_paragraph()
        self._apply_paragraph_style(p, para.style, para.runs)

        for run_data in para.runs:
            run = p.add_run(run_data.text)
            self._apply_run_style(run, run_data.style)

    def _write_table(self, docx, table: Table) -> None:
        """写入表格."""
        if table.rows == 0 or table.cols == 0:
            return

        tbl = docx.add_table(rows=table.rows, cols=table.cols)
        tbl.style = "Table Grid"

        # 应用表格对齐
        if table.style.alignment:
            tbl.alignment = self.ALIGNMENT_MAP.get(table.style.alignment, WD_TABLE_ALIGNMENT.CENTER)
        else:
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_cells in enumerate(table.cells):
            for j, cell_data in enumerate(row_cells):
                if i >= table.rows or j >= table.cols:
                    continue
                cell = tbl.cell(i, j)
                cell.text = ""  # 清除默认内容

                # 写入单元格内容
                para = cell.paragraphs[0]
                if cell_data.runs:
                    for run_data in cell_data.runs:
                        run = para.add_run(run_data.text)
                        self._apply_run_style(run, run_data.style)
                    self._apply_paragraph_style(para, cell_data.style, cell_data.runs)
                else:
                    para.text = cell_data.text

                # 背景色
                if cell_data.shading:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), cell_data.shading.lstrip("#"))
                    cell._element.get_or_add_tcPr().append(shading)

                # 合并单元格
                if cell_data.row_span > 1:
                    for r in range(1, cell_data.row_span):
                        if i + r < table.rows:
                            cell.merge(tbl.cell(i + r, j))
                if cell_data.col_span > 1:
                    for c in range(1, cell_data.col_span):
                        if j + c < table.cols:
                            cell.merge(tbl.cell(i, j + c))

        # 表格标题
        if table.caption:
            cap_para = docx.add_paragraph()
            cap_run = cap_para.add_run(table.caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run.font.bold = True
            cap_run.font.size = Pt(10.5)

    def _write_image(self, docx, img: Image) -> None:
        """写入图片."""
        path = Path(img.path)
        if not path.exists():
            # 图片不存在，添加占位符
            p = docx.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片不存在: {img.path}]")
            run.font.color.rgb = RGBColor(128, 128, 128)
            return

        p = docx.add_paragraph()
        p.alignment = self.ALIGNMENT_MAP.get(img.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        run = p.add_run()
        kwargs = {}
        if img.width:
            kwargs["width"] = Cm(img.width)
        if img.height:
            kwargs["height"] = Cm(img.height)
        run.add_picture(str(path), **kwargs)

        # 图片题注
        if img.caption:
            cap_para = docx.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(img.caption)
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(100, 100, 100)

    def _write_toc_field(self, docx, toc: TOCField) -> None:
        """写入目录域（在 Word 中按 F9 更新即可生成目录）."""
        # 目录标题
        title_para = docx.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(toc.title)
        title_run.font.bold = True
        title_run.font.size = Pt(16)

        # TOC 域代码
        paragraph = docx.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        run2 = paragraph.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f' TOC \\o "1-{toc.max_level}" \\h \\z \\u '
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")
        run3._r.append(fldChar_separate)

        run4 = paragraph.add_run('【请在 Word 中右键此处，选择"更新域"生成目录】')
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run4.font.size = Pt(10)

        run5 = paragraph.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fldChar_end)

    def _write_list(self, docx, lst: DocumentList) -> None:
        """写入列表."""
        for item in lst.items:
            # 使用 Word 的内置列表样式
            style_name = "List Bullet" if lst.list_type.value == "unordered" else "List Number"
            try:
                p = docx.add_paragraph(style=style_name)
            except Exception:
                p = docx.add_paragraph()
                # 手动添加 bullet/number
                prefix = "• " if lst.list_type.value == "unordered" else "1. "
                p.add_run(prefix)

            # 应用样式
            self._apply_paragraph_style(p, item.style, item.runs)

            for run_data in item.runs:
                run = p.add_run(run_data.text)
                self._apply_run_style(run, run_data.style)

    def _write_code_block(self, docx, code: CodeBlock) -> None:
        """写入代码块."""
        for line in code.code.split("\n"):
            p = docx.add_paragraph()
            # 应用代码样式（等宽字体、灰底）
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.15
            pf.left_indent = Pt(24)

            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            # 灰底
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_30

    def _write_blockquote(self, docx, quote: BlockQuote) -> None:
        """写入引用块."""
        p = docx.add_paragraph()
        self._apply_paragraph_style(p, quote.style, quote.runs)

        for run_data in quote.runs:
            run = p.add_run(run_data.text)
            self._apply_run_style(run, run_data.style)

    def _apply_run_style(self, run, style: RunStyle) -> None:
        """应用 Run 样式."""
        if style is None:
            return

        # 字体
        if style.font_name:
            run.font.name = style.font_name
            # 中文字体需要单独设置 eastAsia
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.insert(0, rFonts)
            east_asia = style.font_name_east_asia or style.font_name
            rFonts.set(qn("w:eastAsia"), east_asia)
            rFonts.set(qn("w:ascii"), style.font_name)
            rFonts.set(qn("w:hAnsi"), style.font_name)

        if style.font_size:
            run.font.size = Pt(style.font_size)
        if style.bold is not None:
            run.font.bold = style.bold
        if style.italic is not None:
            run.font.italic = style.italic
        if style.underline is not None:
            run.font.underline = style.underline
        if style.strike is not None:
            run.font.strike = style.strike
        if style.color:
            try:
                run.font.color.rgb = RGBColor.from_string(style.color.lstrip("#"))
            except Exception:
                pass
        if style.highlight:
            try:
                color_map = {
                    "yellow": WD_COLOR_INDEX.YELLOW,
                    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                    "cyan": WD_COLOR_INDEX.TURQUOISE,
                    "magenta": WD_COLOR_INDEX.PINK,
                    "blue": WD_COLOR_INDEX.BLUE,
                    "red": WD_COLOR_INDEX.RED,
                    "gray": WD_COLOR_INDEX.GRAY_50,
                }
                run.font.highlight_color = color_map.get(style.highlight.lower())
            except Exception:
                pass

    def _apply_paragraph_style(self, paragraph, style: ParagraphStyle, runs=None) -> None:
        """应用段落样式."""
        if style is None:
            return

        pf = paragraph.paragraph_format

        if style.alignment:
            paragraph.alignment = self.ALIGNMENT_MAP.get(style.alignment)
        if style.line_spacing is not None:
            pf.line_spacing = style.line_spacing
        if style.space_before is not None:
            pf.space_before = Pt(style.space_before)
        if style.space_after is not None:
            pf.space_after = Pt(style.space_after)

        # 首行缩进：字符数 → 磅值（基于第一个 run 的字号，默认 12pt）
        if style.first_line_indent is not None:
            pf.first_line_indent = Pt(style.first_line_indent)
        elif style.first_line_indent_chars is not None:
            # 按字号计算：2 字符缩进 = 2 * 字号
            font_size = 12  # 默认 12pt
            if runs:
                for r in runs:
                    if r.style.font_size:
                        font_size = r.style.font_size
                        break
            pf.first_line_indent = Pt(style.first_line_indent_chars * font_size)

        if style.left_indent is not None:
            pf.left_indent = Pt(style.left_indent)
        if style.right_indent is not None:
            pf.right_indent = Pt(style.right_indent)
        if style.keep_with_next is not None:
            pf.keep_with_next = style.keep_with_next
        if style.keep_together is not None:
            pf.keep_together = style.keep_together
        if style.page_break_before is not None:
            pf.page_break_before = style.page_break_before
        if style.widow_control is not None:
            pf.widow_control = style.widow_control
