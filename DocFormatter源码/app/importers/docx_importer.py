"""DOCX 导入器 — 将 .docx 文件解析为 Document 模型.

使用 python-docx 读取 Word 文档，提取：
- 段落与标题（按样式名识别）
- Run 级样式（字体、字号、加粗、斜体等）
- 表格（行列结构 + 单元格内容）
- 图片（提取到临时文件，记录路径）
- 页面设置（纸张尺寸、边距）
- 页眉页脚
- 核心属性（标题、作者）

不支持 / 会丢失的内容：
- SmartArt / 复杂文本框 / 艺术字
- 批注 / 修订痕迹
- VBA 宏 / OLE 嵌入对象
- 自定义 XML
"""
import os
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.importers.base import BaseImporter, ImportError
from app.model.document import Document, HeaderFooter
from app.model.elements import (
    Run, Paragraph, Heading, Table, TableCell, Image, PageBreak,
)
from app.model.styles import RunStyle, ParagraphStyle, PageStyle
from app.model.enums import Alignment, HeadingLevel


# Word 对齐 → 枚举 映射
ALIGNMENT_MAP = {
    0: Alignment.LEFT,      # WD_ALIGN_PARAGRAPH.LEFT
    1: Alignment.CENTER,    # WD_ALIGN_PARAGRAPH.CENTER
    2: Alignment.RIGHT,     # WD_ALIGN_PARAGRAPH.RIGHT
    3: Alignment.JUSTIFY,   # WD_ALIGN_PARAGRAPH.JUSTIFY
}


class DocxImporter(BaseImporter):
    """DOCX 文件导入器."""

    @property
    def name(self) -> str:
        return "Word 文档"

    @property
    def supported_extensions(self) -> set:
        return {".docx"}

    def import_file(self, file_path: Path) -> Document:
        self._validate_file(file_path)

        try:
            docx = DocxDocument(str(file_path))
        except Exception as e:
            raise ImportError(f"无法读取 docx 文件: {e}") from e

        doc = Document()
        doc.title = file_path.stem

        # 核心属性
        self._extract_core_properties(docx, doc)

        # 页面设置（取第一个 section）
        if docx.sections:
            self._extract_page_settings(docx.sections[0], doc)
            self._extract_header_footer(docx.sections[0], doc)

        # 文档体：按 body element 顺序遍历
        self._extract_body(docx, doc, file_path)

        return doc

    def _extract_core_properties(self, docx, doc: Document) -> None:
        """提取核心属性."""
        props = docx.core_properties
        if props.title:
            doc.title = props.title
        if props.author:
            doc.author = props.author
        if props.subject:
            doc.subject = props.subject
        if props.keywords:
            doc.keywords = [k.strip() for k in props.keywords.split(",") if k.strip()]

    def _extract_page_settings(self, section, doc: Document) -> None:
        """提取页面设置."""
        ps = doc.page_style

        # 纸张尺寸 (EMU → mm)
        if section.page_width:
            ps.width = section.page_width.mm
        if section.page_height:
            ps.height = section.page_height.mm

        # 边距
        if section.top_margin:
            ps.margin_top = section.top_margin.mm
        if section.bottom_margin:
            ps.margin_bottom = section.bottom_margin.mm
        if section.left_margin:
            ps.margin_left = section.left_margin.mm
        if section.right_margin:
            ps.margin_right = section.right_margin.mm

        # 页眉页脚距离
        if section.header_distance:
            ps.header_distance = section.header_distance.mm
        if section.footer_distance:
            ps.footer_distance = section.footer_distance.mm

        # 方向
        try:
            from docx.enum.section import WD_ORIENT
            if section.orientation == WD_ORIENT.LANDSCAPE:
                ps.orientation = "landscape"
            else:
                ps.orientation = "portrait"
        except Exception:
            pass

    def _extract_header_footer(self, section, doc: Document) -> None:
        """提取页眉页脚."""
        try:
            if section.header and not section.header.is_linked_to_previous:
                hf = self._parse_header_footer(section.header)
                if not hf.is_empty():
                    doc.header = hf

            if section.footer and not section.footer.is_linked_to_previous:
                hf = self._parse_header_footer(section.footer)
                if not hf.is_empty():
                    doc.footer = hf
        except Exception:
            pass  # 页眉页脚解析失败不影响主体

    def _parse_header_footer(self, hf_obj) -> HeaderFooter:
        """解析页眉/页脚对象."""
        result = HeaderFooter()
        for para in hf_obj.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 简化处理：全部放到 center
            result.center = text
        return result

    def _extract_body(self, docx, doc: Document, file_path: Path) -> None:
        """遍历文档体元素."""
        from docx.oxml.ns import qn

        image_dir = self._prepare_image_dir(file_path)

        for element in docx.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落
                para = self._parse_paragraph_element(element, docx, image_dir)
                if para is not None:
                    doc.body.append(para)

            elif tag == "tbl":
                # 表格
                table = self._parse_table_element(element, docx)
                if table is not None:
                    doc.body.append(table)

    def _parse_paragraph_element(self, element, docx, image_dir: Optional[Path]):
        """解析段落 XML 元素."""
        from docx.oxml.ns import qn

        # 查找对应的 Paragraph 对象
        para_obj = None
        for p in docx.paragraphs:
            if p._element is element:
                para_obj = p
                break

        if para_obj is None:
            return None

        # 判断样式：标题还是普通段落
        style_name = para_obj.style.name if para_obj.style else ""
        heading_level = self._detect_heading_level(style_name)

        # 解析 runs
        runs = []
        has_content = False

        for run_obj in para_obj.runs:
            run = self._parse_run(run_obj, image_dir)
            if run is not None:
                runs.append(run)
                if run.text.strip() or hasattr(run, '_image_path'):
                    has_content = True

        # 检查是否有图片 run（通过 _image_path 标记）
        has_image = any(hasattr(r, '_image_path') for r in runs)

        # 检查是否有分页符
        has_page_break = self._has_page_break(element)

        # 构建元素
        if heading_level is not None:
            result = Heading(
                runs=runs,
                level=HeadingLevel(heading_level),
                style=self._parse_paragraph_format(para_obj),
            )
        else:
            result = Paragraph(
                runs=runs,
                style=self._parse_paragraph_format(para_obj),
                style_name=style_name if style_name else None,
            )

        # 如果是纯分页符段落，返回 PageBreak
        if has_page_break and not has_content and not has_image:
            return PageBreak()

        return result if has_content or has_image or has_page_break else None

    def _parse_run(self, run_obj, image_dir: Optional[Path]) -> Optional[Run]:
        """解析单个 Run."""
        from docx.oxml.ns import qn

        text = run_obj.text or ""
        font = run_obj.font

        style = RunStyle(
            font_name=font.name,
            font_size=font.size.pt if font.size else None,
            bold=font.bold,
            italic=font.italic,
            underline=font.underline,
            strike=font.strike,
        )

        # 中文字体
        try:
            rpr = run_obj._element.find(qn("w:rPr"))
            if rpr is not None:
                rFonts = rpr.find(qn("w:rFonts"))
                if rFonts is not None:
                    east_asia = rFonts.get(qn("w:eastAsia"))
                    if east_asia:
                        style.font_name_east_asia = east_asia
        except Exception:
            pass

        # 颜色
        if font.color and font.color.rgb:
            style.color = f"#{font.color.rgb}"

        run = Run(text=text, style=style)

        # 检查 run 中是否有图片
        drawings = run_obj._element.findall(qn("w:drawing"))
        for drawing in drawings:
            blip = drawing.find(".//" + qn("a:blip"))
            if blip is not None:
                embed_id = blip.get(qn("r:embed"))
                if embed_id:
                    try:
                        image_path = self._extract_image(
                            run_obj.part, embed_id, image_dir
                        )
                        if image_path:
                            # 把图片路径附加到 run 上，后续处理
                            run._image_path = image_path
                    except Exception:
                        pass

        return run

    def _parse_paragraph_format(self, para_obj) -> ParagraphStyle:
        """解析段落格式."""
        pf = para_obj.paragraph_format
        style = ParagraphStyle()

        # 对齐
        if para_obj.alignment is not None:
            style.alignment = ALIGNMENT_MAP.get(int(para_obj.alignment))

        # 行距
        if pf.line_spacing is not None:
            style.line_spacing = pf.line_spacing

        # 段前段后
        if pf.space_before is not None:
            style.space_before = pf.space_before.pt
        if pf.space_after is not None:
            style.space_after = pf.space_after.pt

        # 缩进
        if pf.first_line_indent is not None:
            style.first_line_indent = pf.first_line_indent.pt
        if pf.left_indent is not None:
            style.left_indent = pf.left_indent.pt
        if pf.right_indent is not None:
            style.right_indent = pf.right_indent.pt

        # 其他
        if pf.keep_with_next is not None:
            style.keep_with_next = pf.keep_with_next
        if pf.keep_together is not None:
            style.keep_together = pf.keep_together
        if pf.page_break_before is not None:
            style.page_break_before = pf.page_break_before

        return style

    def _parse_table_element(self, element, docx) -> Optional[Table]:
        """解析表格 XML 元素."""
        from docx.oxml.ns import qn

        # 查找对应的 Table 对象
        table_obj = None
        for t in docx.tables:
            if t._element is element:
                table_obj = t
                break

        if table_obj is None:
            return None

        rows = len(table_obj.rows)
        cols = len(table_obj.columns)
        cells = []

        for row in table_obj.rows:
            row_cells = []
            for cell in row.cells:
                text = cell.text.strip()
                para_style = ParagraphStyle()
                if cell.paragraphs:
                    para_style = self._parse_paragraph_format(cell.paragraphs[0])
                row_cells.append(TableCell(
                    text=text,
                    style=para_style,
                ))
            cells.append(row_cells)

        return Table(rows=rows, cols=cols, cells=cells)

    def _detect_heading_level(self, style_name: str) -> Optional[int]:
        """从样式名推断标题层级."""
        if not style_name:
            return None
        style_lower = style_name.lower()
        if style_lower.startswith("heading"):
            try:
                level = int(style_lower.replace("heading", "").strip())
                if 1 <= level <= 4:
                    return level
            except ValueError:
                pass
        if style_lower == "title":
            return 1
        return None

    def _has_page_break(self, element) -> bool:
        """检查段落是否包含分页符."""
        from docx.oxml.ns import qn
        # 查找 lastRenderedPageBreak 或 br[@w:type='page']
        for br in element.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    def _prepare_image_dir(self, file_path: Path) -> Optional[Path]:
        """准备图片提取目录."""
        try:
            image_dir = file_path.parent / f".{file_path.stem}_images"
            image_dir.mkdir(exist_ok=True)
            return image_dir
        except Exception:
            return None

    def _extract_image(self, part, embed_id: str, image_dir: Optional[Path]) -> Optional[str]:
        """提取嵌入的图片到临时目录."""
        if image_dir is None:
            return None
        try:
            rel = part.rels[embed_id]
            image_part = rel.target_part
            # 获取图片扩展名
            content_type = image_part.content_type
            ext_map = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/gif": ".gif",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
                "image/x-emf": ".emf",
                "image/x-wmf": ".wmf",
            }
            ext = ext_map.get(content_type, ".png")

            # 生成文件名
            image_name = f"{embed_id}{ext}"
            image_path = image_dir / image_name

            # 写入文件
            with open(image_path, "wb") as f:
                f.write(image_part.blob)

            return str(image_path)
        except Exception:
            return None
